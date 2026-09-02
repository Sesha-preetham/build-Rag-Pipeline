"""Generate the three sample documents used by the W8 notebooks.

Reproducible — learners can rerun this to regenerate the sample docs, or
adapt it to generate more variety.

Alternatively, learners can substitute their own public-domain samples:
  - PDF:  https://www.pdf-tools.com/pdf/sample.pdf   (any short public PDF works)
  - HTML: save any Wikipedia article as an HTML file
  - DOCX: any DOCX with heading styles (Word's default templates are fine)

Just place them at demos/sample_docs/ with the expected filenames.

All fake PII in these docs is invented for demo purposes.
"""
from pathlib import Path
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt

OUT = Path(__file__).parent / "sample_docs"
OUT.mkdir(exist_ok=True)


# ─── 1. PDF: company_policies.pdf ────────────────────────────────────
# 3 pages with headings + one table + fake PII in HR contact section.

def make_pdf():
    path = OUT / "company_policies.pdf"
    doc = SimpleDocTemplate(str(path), pagesize=LETTER,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    heading = ParagraphStyle('h', parent=styles['Heading1'], fontSize=16, spaceAfter=12)
    subhead = ParagraphStyle('h2', parent=styles['Heading2'], fontSize=13, spaceAfter=8)
    body = ParagraphStyle('body', parent=styles['BodyText'], fontSize=10, spaceAfter=8, leading=14)
    
    story = []
    
    # Page 1 — Leave Policy
    story.append(Paragraph("Acme Corp — Employee Handbook", heading))
    story.append(Paragraph("Section 1: Leave Policy", subhead))
    story.append(Paragraph(
        "Employees are entitled to 20 days of paid annual leave per calendar year. "
        "Leave must be requested at least two weeks in advance through the HR portal. "
        "Unused leave may be carried over up to a maximum of 5 days into the following year.",
        body))
    story.append(Paragraph(
        "Sick leave is separate and unlimited, subject to a doctor's note for absences "
        "longer than three consecutive working days. Bereavement leave of up to 5 days "
        "is granted for immediate family members.",
        body))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Section 2: Leave Approval Table", subhead))
    
    # A simple table — pdfplumber handles this better than PyMuPDF
    table_data = [
        ["Leave Type",  "Days per year", "Approval Needed"],
        ["Annual",      "20",            "Line manager"],
        ["Sick",        "Unlimited",     "None (< 3 days)"],
        ["Bereavement", "5",             "HR notification"],
        ["Parental",    "12 weeks",      "HR + Legal"],
    ]
    t = Table(table_data, colWidths=[2*inch, 1.5*inch, 2*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0D243D')),
        ('TEXTCOLOR',  (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN',      (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('GRID',       (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    story.append(t)
    story.append(PageBreak())
    
    # Page 2 — Remote Work
    story.append(Paragraph("Section 3: Remote Work Policy", subhead))
    story.append(Paragraph(
        "All full-time employees are eligible for hybrid work arrangements. The standard "
        "expectation is three days in-office per week, though team leads may adjust this "
        "based on project needs. Full-remote status requires VP approval and a documented "
        "business case.",
        body))
    story.append(Paragraph(
        "Home office equipment stipends of up to $500 are available annually for employees "
        "working remotely at least two days per week. Submit receipts through the expense "
        "portal within 30 days of purchase.",
        body))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Section 4: Working Hours", subhead))
    story.append(Paragraph(
        "Core hours are 10am to 4pm local time. Employees are expected to be reachable "
        "during core hours regardless of remote or in-office status. Flexible start and "
        "end times are supported provided daily hours total 8.",
        body))
    story.append(PageBreak())
    
    # Page 3 — HR Contact (with fake PII deliberately embedded)
    story.append(Paragraph("Section 5: HR Contact Information", subhead))
    story.append(Paragraph(
        "For questions about this handbook, contact the HR team. Our HR Business Partner "
        "for Engineering is Sarah Chen — she can be reached at sarah.chen@acme.com or "
        "on her direct line at +1-555-0142. For urgent after-hours matters, the on-call "
        "HR number is +1-555-0199.",
        body))
    story.append(Paragraph(
        "Employees can also submit anonymous concerns through the ethics hotline at "
        "ethics@acme-hotline.com. All submissions are handled by our external partner "
        "and are not visible to internal HR staff.",
        body))
    story.append(Paragraph(
        "Payroll queries should be directed to Michael Torres in Finance, employee ID "
        "EMP-04521, at michael.torres@acme.com or extension 4521.",
        body))
    
    doc.build(story)
    print(f"  ✓ {path.name} ({path.stat().st_size} bytes)")


# ─── 2. HTML: product_page.html ──────────────────────────────────────
# Realistic marketing page with nav, main content, footer. No PII.

def make_html():
    path = OUT / "product_page.html"
    html = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>Acme Analytics Platform — Product Overview</title>
  <script>
    // Analytics tracking — should be stripped by BeautifulSoup
    window.trackPageView('product_overview');
  </script>
  <style>body { font-family: Arial, sans-serif; }</style>
</head>
<body>
  <nav id="site-nav">
    <ul>
      <li><a href="/">Home</a></li>
      <li><a href="/products">Products</a></li>
      <li><a href="/pricing">Pricing</a></li>
      <li><a href="/docs">Docs</a></li>
    </ul>
  </nav>

  <main>
    <article>
      <h1>Acme Analytics Platform</h1>
      <p class="lead">
        Turn raw event data into actionable dashboards without writing SQL. The Acme
        Analytics Platform sits on top of your data warehouse and gives every team
        member a drag-and-drop interface for building charts, reports, and alerts.
      </p>

      <h2>Key Features</h2>
      <ul>
        <li>Native connectors for Snowflake, BigQuery, Redshift, and Postgres</li>
        <li>Real-time streaming ingestion with sub-second latency</li>
        <li>Role-based access control with row-level security policies</li>
        <li>Embedded analytics — drop dashboards into your own product</li>
      </ul>

      <h2>How It Works</h2>
      <p>
        Point the platform at your data warehouse. Our schema crawler catalogs
        tables and infers relationships. Business users then build queries using
        the visual interface. Behind the scenes, the platform compiles those
        queries to optimised SQL and caches results at multiple layers.
      </p>
      <p>
        For advanced use cases, engineers can drop into a code editor and write
        Python or SQL directly. The platform supports version control, code
        review, and CI/CD for analytics code.
      </p>

      <h2>Pricing</h2>
      <p>
        The platform is priced per active user per month, with volume discounts
        starting at 50 seats. Enterprise contracts include SSO, audit logging,
        and dedicated support. See the <a href="/pricing">pricing page</a> for
        current rates.
      </p>
    </article>
  </main>

  <footer id="site-footer">
    <p>© 2025 Acme Corp. All rights reserved.</p>
    <ul>
      <li><a href="/privacy">Privacy</a></li>
      <li><a href="/terms">Terms</a></li>
      <li><a href="/contact">Contact</a></li>
    </ul>
  </footer>
</body>
</html>
"""
    path.write_text(html)
    print(f"  ✓ {path.name} ({path.stat().st_size} bytes)")


# ─── 3. DOCX: onboarding.docx ────────────────────────────────────────
# DOCX with heading styles so python-docx can extract structure.
# Also has fake PII deliberately embedded.

def make_docx():
    path = OUT / "onboarding.docx"
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    doc.add_heading('New Employee Onboarding Guide', level=0)
    
    doc.add_heading('Week 1: Getting Set Up', level=1)
    doc.add_paragraph(
        "Welcome to Acme Corp. Your first week focuses on getting your accounts "
        "provisioned and meeting your immediate team. Your manager will schedule "
        "1:1 meetings with each team member and walk you through current projects."
    )
    doc.add_paragraph(
        "Before your first day, IT will have provisioned your laptop, email account, "
        "and access to core systems (Slack, GitHub, Jira, Google Workspace). If any "
        "credential is missing on day 1, contact IT support immediately."
    )
    
    doc.add_heading('IT Support Contacts', level=2)
    # PII deliberately embedded here for the Day 2 scrubbing demo
    doc.add_paragraph(
        "For laptop or account issues, contact IT support at it-help@acme.com or "
        "call +1-555-0180 during business hours. Our IT lead is Priya Sharma "
        "(priya.sharma@acme.com, employee ID EMP-01847)."
    )
    
    doc.add_heading('Week 2: Team Immersion', level=1)
    doc.add_paragraph(
        "By week 2, you should be attending your team's regular meetings including "
        "sprint planning, daily standups, and retros. Your manager will pair you with "
        "a peer buddy who can answer day-to-day questions."
    )
    doc.add_paragraph(
        "This week is a good time to start pairing on real work. Small, well-scoped "
        "tasks help you learn the codebase without blocking on unknowns. Your buddy "
        "or manager will help identify appropriate starter tasks."
    )
    
    doc.add_heading('Week 3-4: Owning Your First Feature', level=1)
    doc.add_paragraph(
        "By the end of your first month, you should own at least one small feature or "
        "improvement end-to-end — from design through code review to production deploy. "
        "This is the practical goal of onboarding: shipped code, not just knowledge."
    )
    
    doc.add_heading('30-Day Check-In', level=2)
    doc.add_paragraph(
        "At the 30-day mark, your manager will schedule a formal check-in to discuss "
        "how the onboarding is going, gather your feedback, and adjust expectations "
        "for the next 60 days. Prepare 2-3 things that are going well and 2-3 things "
        "that are blocking you."
    )
    
    doc.save(str(path))
    print(f"  ✓ {path.name} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    print(f"Generating sample documents in {OUT}...")
    make_pdf()
    make_html()
    make_docx()
    print("Done.")
    print("")
    print("Want more variety? Two options:")
    print("  1. Modify this script — add more sections, more heading levels, more PII patterns")
    print("  2. Replace with public-domain samples — any short PDF/HTML/DOCX works if")
    print("     placed at demos/sample_docs/ with the expected filenames:")
    print("       company_policies.pdf, product_page.html, onboarding.docx")
